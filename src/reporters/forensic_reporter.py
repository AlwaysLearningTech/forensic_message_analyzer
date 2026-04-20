import pandas as pd
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any
import logging
import json
import re
import html as html_module
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from ..config import Config
from ..forensic_utils import ForensicRecorder
from ..utils.legal_compliance import LegalComplianceManager
from ..utils.pricing import get_pricing
from .report_utils import match_quote_to_message, generate_limitations, markdown_to_docx

logger = logging.getLogger(__name__)


class ForensicReporter:
    """
    Generate forensic reports in multiple formats.
    Ensures legal defensibility and chain of custody documentation.
    """

    def __init__(self, forensic_recorder: ForensicRecorder, config: Config = None):
        """
        Initialize the forensic reporter.

        Args:
            forensic_recorder: ForensicRecorder instance for chain of custody
            config: Config instance. If None, creates a new one.
        """
        self.config = config if config is not None else Config()
        self.forensic = forensic_recorder
        self.compliance = LegalComplianceManager(config=self.config, forensic_recorder=forensic_recorder)
        self.output_dir = Path(self.config.output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # Shared helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _compute_date_range(messages) -> str:
        """Return a 'YYYY-MM-DD to YYYY-MM-DD' string from a list of message dicts."""
        if not messages:
            return 'N/A'
        timestamps = [msg.get('timestamp') for msg in messages if msg.get('timestamp') is not None]
        if not timestamps:
            return 'N/A'
        dt_timestamps = []
        for ts in timestamps:
            try:
                if isinstance(ts, str):
                    parsed = pd.to_datetime(ts, utc=True)
                    if not pd.isna(parsed):
                        dt_timestamps.append(parsed)
                elif hasattr(ts, 'year') and not pd.isna(ts):
                    if hasattr(ts, 'tzinfo') and ts.tzinfo is None:
                        ts = ts.replace(tzinfo=pd.Timestamp.now(tz='UTC').tzinfo)
                    dt_timestamps.append(ts)
            except Exception:
                pass
        if not dt_timestamps:
            return 'N/A'
        return f"{min(dt_timestamps).strftime('%Y-%m-%d')} to {max(dt_timestamps).strftime('%Y-%m-%d')}"

    @staticmethod
    def _esc(text) -> str:
        """Escape text for safe use in ReportLab Paragraph (XML/HTML context)."""
        return html_module.escape(str(text)) if text else ''

    @staticmethod
    def _match_quote_to_message(quote: str, messages: list) -> dict:
        """Match an AI-identified quote to its source message via substring matching."""
        return match_quote_to_message(quote, messages)

    def _docx_to_pdf(self, docx_path: Path) -> Path:
        """Convert a DOCX file to PDF using docx2pdf (MS Word / LibreOffice).

        On macOS, MS Word is sandboxed and may lack permission to read/write arbitrary directories. We work around this by copying the DOCX into a temporary directory under ~/Documents (which Word always has access to), converting there, then moving the PDF back to the original location.

        Returns the path to the generated PDF file.
        Raises RuntimeError if the conversion fails.
        """
        import shutil
        import tempfile
        from docx2pdf import convert

        pdf_path = docx_path.with_suffix('.pdf')

        # Use ~/Documents as the temp root — Word always has access to it.
        docs_dir = Path.home() / "Documents"
        tmp_dir = Path(tempfile.mkdtemp(dir=docs_dir if docs_dir.is_dir() else None,
                                        prefix='.forensic_pdf_'))
        try:
            tmp_docx = tmp_dir / docx_path.name
            tmp_pdf = tmp_dir / pdf_path.name
            shutil.copy2(docx_path, tmp_docx)
            convert(str(tmp_docx), str(tmp_pdf))
            if not tmp_pdf.exists():
                raise RuntimeError(f"MS Word failed to produce {tmp_pdf.name} — Word may need to be restarted or granted Full Disk Access in System Settings > Privacy")
            shutil.move(str(tmp_pdf), str(pdf_path))
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

        file_hash = self.forensic.compute_hash(pdf_path)
        self.forensic.record_action(
            "pdf_converted",
            f"Converted {docx_path.name} to PDF with hash {file_hash}",
            {"source": str(docx_path), "path": str(pdf_path), "hash": file_hash}
        )
        return pdf_path

    @staticmethod
    def _style_docx_table(table) -> None:
        """Apply consistent Microsoft blue theme styling to a python-docx table.

        Header: dark blue (#1F4E79) with white bold text.
        Body: alternating white / light blue (#D6E4F0).
        Borders: thin grey.
        """
        HEADER_BG = '1F4E79'
        ALT_BG = 'D6E4F0'

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                # Set background
                if row_idx == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG}"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(10)
                elif row_idx % 2 == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ALT_BG}"/>')
                    cell._tc.get_or_add_tcPr().append(shading)

                # Set font size for body rows
                if row_idx > 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    @staticmethod
    def _render_methodology_to_docx(doc, sections, base_level: int = 1) -> None:
        """Render structured methodology sections into a python-docx document."""
        for section in sections:
            doc.add_heading(section['heading'], level=base_level)
            for block in section['blocks']:
                btype = block['type']
                if btype == 'paragraph':
                    doc.add_paragraph(block['text'])
                elif btype == 'bullets':
                    for item in block['items']:
                        doc.add_paragraph(item, style='List Bullet')
                elif btype == 'definition':
                    para = doc.add_paragraph()
                    run = para.add_run(f"{block['term']}. ")
                    run.bold = True
                    para.add_run(block['text'])

    def generate_comprehensive_report(self,
                                     extracted_data: Dict,
                                     analysis_results: Dict,
                                     review_decisions: Dict) -> Dict[str, Path]:
        """
        Generate comprehensive forensic report in multiple formats.
        
        Args:
            extracted_data: Data from extraction phase
            analysis_results: Results from analysis phase
            review_decisions: Manual review decisions
            
        Returns:
            Dictionary mapping format to output file path
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        reports = {}

        # Generate legal team summary first (used in Word/PDF reports)
        legal_summary = self._generate_legal_team_summary(
            extracted_data, analysis_results, review_decisions
        )

        # Generate Word report
        try:
            word_path = self._generate_word_report(
                extracted_data, analysis_results, review_decisions, timestamp,
                legal_summary=legal_summary
            )
            reports['word'] = word_path
            logger.info(f"Generated Word report: {word_path}")
        except Exception as e:
            import traceback
            logger.error(f"Failed to generate Word report: {e}")
            logger.error(traceback.format_exc())
            self.forensic.record_action(
                "report_generation_error",
                f"Word report generation failed: {str(e)}"
            )

        # Generate standalone Methodology document (lay-friendly, distinct from the findings report so the legal team can read it without wading through case-specific results)
        try:
            methodology_path = self._generate_methodology_document(
                extracted_data, timestamp
            )
            reports['methodology'] = methodology_path
            logger.info(f"Generated Methodology document: {methodology_path}")
        except Exception as e:
            logger.error(f"Failed to generate methodology document: {e}")
            self.forensic.record_action(
                "report_generation_error",
                f"Methodology document generation failed: {str(e)}"
            )

        # PDF versions: convert each DOCX to PDF via docx2pdf for exact fidelity
        if 'methodology' in reports:
            try:
                methodology_pdf = self._docx_to_pdf(reports['methodology'])
                reports['methodology_pdf'] = methodology_pdf
                logger.info(f"Generated Methodology PDF: {methodology_pdf}")
            except Exception as e:
                logger.error(f"Failed to convert methodology to PDF: {e}")
                self.forensic.record_action(
                    "report_generation_error",
                    f"Methodology PDF conversion failed: {str(e)}"
                )

        if 'word' in reports:
            try:
                pdf_path = self._docx_to_pdf(reports['word'])
                reports['pdf'] = pdf_path
                logger.info(f"Generated PDF report: {pdf_path}")
            except Exception as e:
                logger.error(f"Failed to convert Word report to PDF: {e}")
                self.forensic.record_action(
                    "report_generation_error",
                    f"PDF report conversion failed: {str(e)}"
                )
        
        # Generate JSON report
        try:
            json_path = self._generate_json_report(
                extracted_data, analysis_results, review_decisions, timestamp,
                legal_summary=legal_summary
            )
            reports['json'] = json_path
            logger.info(f"Generated JSON report: {json_path}")
        except Exception as e:
            logger.error(f"Failed to generate JSON report: {e}")
            self.forensic.record_action(
                "report_generation_error",
                f"JSON report generation failed: {str(e)}"
            )

        # Store legal summary text for deferred docx generation (after all reports exist)
        self._legal_summary_text = legal_summary

        # Record report generation
        self.forensic.record_action(
            "reports_generated",
            f"Generated {len(reports)} forensic reports",
            {
                "formats": list(reports.keys()),
                "timestamp": timestamp
            }
        )
        
        return reports
    
    def _generate_methodology_document(self, extracted_data: Dict, timestamp: str) -> Path:
        """Generate a standalone Methodology Statement Word document.

        Separate from the findings report so the legal team (and the court) can read the methodology without having to navigate case-specific results. Contents are produced by LegalComplianceManager.generate_methodology_sections(), which is plain-language and tied to FRE / Daubert factors point by point.
        """
        doc = Document()

        # Title
        title = doc.add_heading('Methodology Statement', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Case header
        header = self.compliance.generate_report_header()
        case_numbers = header.get('case_numbers') or [header['case_number']]
        if len(case_numbers) > 1:
            doc.add_paragraph('Case Numbers:')
            for cn in case_numbers:
                doc.add_paragraph(cn, style='List Bullet')
        else:
            doc.add_paragraph(f"Case Number: {case_numbers[0]}")
        if header['case_name'] != 'Not assigned':
            doc.add_paragraph(f"Case Name: {header['case_name']}")
        doc.add_paragraph(f"Generated: {header['date_of_examination']}")
        doc.add_paragraph('')

        # Methodology body — structured sections render as real headings
        sections = self.compliance.generate_methodology_sections()
        self._render_methodology_to_docx(doc, sections, base_level=1)

        # Standards compliance — rendered through the same structured renderer as the methodology body so headings, bullet lists, and term/definition pairs survive instead of coming out as a flat text block.
        doc.add_page_break()
        doc.add_heading('Standards Compliance', level=0)
        standards_sections = self.compliance.generate_standards_compliance_sections()
        self._render_methodology_to_docx(doc, standards_sections, base_level=1)

        # Completeness Validation
        messages = extracted_data.get('messages', extracted_data.get('combined', []))
        completeness = self.compliance.validate_completeness(messages)
        doc.add_page_break()
        doc.add_heading('Completeness Validation (FRE 106)', level=1)
        doc.add_paragraph(
            f"Total messages examined: {completeness.get('total_messages', 0)}. "
            f"Conversations analysed: {len(completeness.get('conversations', {}))}. "
            f"Complete: {'Yes' if completeness.get('is_complete') else 'No'}."
        )
        issues = completeness.get('issues', [])
        if issues:
            doc.add_paragraph('Issues detected (review and supplement as needed):')
            for issue in issues:
                doc.add_paragraph(issue, style='List Bullet')
        else:
            doc.add_paragraph('No completeness issues detected.')

        output_path = self.output_dir / f"methodology_{timestamp}.docx"
        doc.save(output_path)

        file_hash = self.forensic.compute_hash(output_path)
        self.forensic.record_action(
            "methodology_document_generated",
            f"Generated standalone methodology document with hash {file_hash}",
            {"path": str(output_path), "hash": file_hash}
        )
        return output_path

    def _build_cover_sheet_content(self, reports: Dict[str, Any], timestamp: str) -> Dict[str, Any]:
        """Return the structured content of the READ ME FIRST cover sheet.

        Shared by the DOCX and PDF renderers so both formats emit identical text. Returns a dict with ``header`` (list of label/value pairs), ``intro`` (paragraph), ``guide`` (list of (question, filename, description) tuples), and ``footer`` (italic paragraph).
        """
        def _name(key: str) -> str:
            value = reports.get(key)
            return Path(str(value)).name if value else ''

        methodology_name = _name('methodology')
        legal_summary_name = _name('legal_summary')
        full_report_name = _name('pdf') or _name('word')
        chat_name = _name('chat') or _name('chat_html')
        events_timeline_name = _name('events_timeline')
        timeline_name = _name('timeline')
        html_name = _name('html')
        excel_name = _name('excel')

        header_meta = self.compliance.generate_report_header()
        case_numbers = header_meta.get('case_numbers') or [header_meta['case_number']]
        header_rows = [
            ('Case Number(s)', '; '.join(case_numbers)),
        ]
        if header_meta['case_name'] and header_meta['case_name'] != 'Not assigned':
            header_rows.append(('Case Name', header_meta['case_name']))
        header_rows.append(('Generated', header_meta['date_of_examination']))
        if header_meta.get('examiner_name') and header_meta['examiner_name'] != 'Not specified':
            header_rows.append(('Examiner', header_meta['examiner_name']))

        guide: list = []
        if methodology_name:
            guide.append((
                'If anyone challenges the methods or the science',
                methodology_name,
                'Plain-language, judge-readable walkthrough of every step the analyzer took, with an explicit point-by-point map of how each Federal Rule of Evidence and Daubert factor was satisfied. Includes empirical citations for every pattern used to flag a message. Read this first if methodology is questioned.'
            ))
        if legal_summary_name:
            guide.append((
                'If you want the plain-English findings',
                legal_summary_name,
                'AI-assisted narrative summary written for attorneys: what was found, what it appears to mean, what to do next, and a guide to the rest of the files in this package.'
            ))
        if full_report_name:
            guide.append((
                'If you want the full record for filing or distribution',
                full_report_name,
                'Comprehensive forensic report: case information, findings summary, threat analysis, sentiment analysis, manual-review breakdown, and chain-of-custody reference. The authoritative document for the case file.'
            ))
        if events_timeline_name:
            guide.append((
                'If you want the big-picture chronology of the case',
                events_timeline_name,
                'Sparse, court-facing timeline showing only reviewer-confirmed events — confirmed threats, coercive-control pattern clusters, and tone shifts — with category badges. Open in a web browser.'
            ))
        if chat_name:
            guide.append((
                'If you want to read the conversations themselves',
                chat_name,
                'iMessage-style chat-bubble HTML transcript of the relevant conversations, with inline images, edit history, and deletion / URL-preview / shared-location markers. Open in a web browser.'
            ))
        if timeline_name:
            guide.append((
                'If you want a minute-level message-by-message timeline',
                timeline_name,
                'Detailed chronological view with every flagged event and all email communications; for analyst drill-down rather than legal-team reading.'
            ))
        if excel_name:
            guide.append((
                'If you want to sort, filter, or query the data yourself',
                excel_name,
                'Multi-sheet Excel workbook: per-person message tabs, findings summary, timeline, conversation threads, manual-review decisions, and third-party contacts.'
            ))
        if html_name:
            guide.append((
                'If you want a printable visual report with attachments',
                html_name,
                'HTML report with inline base64 attachment images and the three legal appendices (Methodology, Completeness Validation, Limitations).'
            ))
        guide.append((
            'If you need the technical audit trail (for a forensics expert)',
            f'chain_of_custody_{timestamp}.json',
            'Timestamped audit trail of every operation performed during the run, with SHA-256 hashes of every input and output file. This is for a digital-forensics expert; the methodology document above is what to give a judge or attorney.'
        ))

        return {
            'title': 'READ ME FIRST',
            'subtitle': 'Forensic Analysis Report Package — Reading Guide',
            'header_rows': header_rows,
            'intro': (
                'This package contains several documents. Each one answers a different question. '
                'Open the document below that matches what you need; every document references the '
                'others by filename so you can navigate between them.'
            ),
            'guide': guide,
            'footer': (
                'All files in this package were produced by the same analysis run and are forensically '
                'linked through SHA-256 hashes recorded in the chain of custody. Every artifact is signed '
                'with a detached Ed25519 signature; see the accompanying .sig and .sig.pub files.'
            ),
        }

    def generate_cover_sheet(self, reports: Dict[str, Any], timestamp: str) -> Dict[str, Path]:
        """Generate the READ ME FIRST cover sheet in both DOCX and PDF form.

        Args:
            reports: Mapping of report-format keys to file paths.
            timestamp: Run timestamp used for the filenames.

        Returns:
            Dict with ``docx`` and ``pdf`` keys pointing at the two written files.
        """
        content = self._build_cover_sheet_content(reports, timestamp)
        docx_path = self._render_cover_sheet_docx(content, timestamp)
        pdf_path = self._docx_to_pdf(docx_path)
        return {'docx': docx_path, 'pdf': pdf_path}

    def _render_cover_sheet_docx(self, content: Dict[str, Any], timestamp: str) -> Path:
        """Render the cover-sheet content dict to a Word document."""

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        title = doc.add_heading(content['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(content['subtitle'])
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(12)

        for label, value in content['header_rows']:
            line = doc.add_paragraph()
            line.add_run(f'{label}: ').bold = True
            line.add_run(str(value))

        doc.add_paragraph(content['intro'])
        doc.add_heading('Where to start', level=1)

        for question, filename, description in content['guide']:
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f'{question}:')
            q_run.bold = True

            f_para = doc.add_paragraph()
            f_para.paragraph_format.left_indent = Inches(0.25)
            f_run = f_para.add_run(f'→ Open  {filename}')
            f_run.font.name = 'Consolas'
            f_run.font.size = Pt(10)

            d_para = doc.add_paragraph()
            d_para.paragraph_format.left_indent = Inches(0.25)
            d_para.paragraph_format.space_after = Pt(6)
            d_run = d_para.add_run(description)
            d_run.font.size = Pt(10)

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run(content['footer'])
        footer_run.italic = True
        footer_run.font.size = Pt(9)

        output_path = self.output_dir / f"READ_ME_FIRST_{timestamp}.docx"
        doc.save(output_path)

        file_hash = self.forensic.compute_hash(output_path)
        self.forensic.record_action(
            "cover_sheet_generated",
            f"Generated READ ME FIRST cover sheet (docx) with hash {file_hash}",
            {"path": str(output_path), "hash": file_hash}
        )
        return output_path

    def _generate_word_report(self, extracted_data: Dict, analysis_results: Dict,
                            review_decisions: Dict, timestamp: str,
                            legal_summary: str = None) -> Path:
        """Generate Word document report."""
        doc = Document()
        
        # Title page
        title = doc.add_heading('Forensic Message Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Generated: {self.compliance.format_timestamp()}')
        doc.add_paragraph(f'Case ID: {timestamp}')
        doc.add_page_break()

        # ----- Legal Compliance Header -----
        header = self.compliance.generate_report_header()
        doc.add_heading('Case Information', 1)
        case_numbers = header.get('case_numbers') or [header['case_number']]
        case_info_rows = [
            ('Field', 'Value'),
            ('Case Number(s)', '\n'.join(str(cn) for cn in case_numbers)),
            ('Case Name', header['case_name']),
            ('Examiner', header['examiner_name']),
            ('Organization', header['organization']),
            ('Date of Examination', header['date_of_examination']),
            ('Tools Used', header['tools_used']),
        ]
        table = doc.add_table(rows=len(case_info_rows), cols=2)
        for row_idx, (field, value) in enumerate(case_info_rows):
            table.rows[row_idx].cells[0].text = field
            table.rows[row_idx].cells[1].text = str(value)
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(4.0)
        self._style_docx_table(table)
        doc.add_paragraph('')  # spacer

        # Methodology Statement
        doc.add_heading('Methodology', 1)
        sections = self.compliance.generate_methodology_sections()
        self._render_methodology_to_docx(doc, sections, base_level=2)

        # Standards Compliance Statement
        doc.add_heading('Standards Compliance', 1)
        standards_sections = self.compliance.generate_standards_compliance_sections()
        self._render_methodology_to_docx(doc, standards_sections, base_level=2)

        # Completeness Validation (FRE 106)
        messages = extracted_data.get('messages', extracted_data.get('combined', []))
        completeness = self.compliance.validate_completeness(messages)
        doc.add_heading('Completeness Validation', 1)
        doc.add_paragraph(
            f"Total messages: {completeness.get('total_messages', 0)}. "
            f"Conversations analyzed: {len(completeness.get('conversations', {}))}. "
            f"Complete: {'Yes' if completeness.get('is_complete') else 'No'}."
        )
        issues = completeness.get('issues', [])
        if issues:
            doc.add_paragraph('Issues detected:')
            for issue in issues:
                doc.add_paragraph(issue, style='List Bullet')

        # Limitations
        doc.add_heading('Limitations', 1)
        limitations = self._generate_limitations(analysis_results)
        for item in limitations:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_page_break()

        # === AI-Powered Findings Summary ===
        ai_analysis = analysis_results.get('ai_analysis', {})
        if ai_analysis and ai_analysis.get('conversation_summary') and \
           'not configured' not in ai_analysis.get('conversation_summary', '').lower():
            doc.add_heading('Findings Summary', 1)
            doc.add_paragraph(
                'This section consolidates the analysis findings for rapid legal team review. '
                'All flagged items — regardless of whether they were surfaced by pattern '
                'matching, statistical analysis, or AI — were submitted to the same manual '
                'review process; only items confirmed during manual review are reflected '
                'in the threat counts below.'
            )

            # Executive Summary
            doc.add_heading('Analysis Overview', 2)
            doc.add_paragraph(ai_analysis.get('conversation_summary', 'Not available'))

            # Risk indicators with severity
            risk_indicators = ai_analysis.get('risk_indicators', [])
            if risk_indicators:
                doc.add_heading('Risk Indicators', 2)
                for risk in risk_indicators:
                    if isinstance(risk, dict):
                        severity = str(risk.get('severity', 'unknown')).upper()
                        indicator = risk.get('indicator', risk.get('description', risk.get('detail', '')))
                        action = risk.get('recommended_action', '')
                        doc.add_paragraph(f'[{severity}] {indicator}')
                        if action:
                            doc.add_paragraph(f'    Recommended: {action}')
                    else:
                        doc.add_paragraph(f'  {risk}')

            # Notable quotes
            notable_quotes = ai_analysis.get('notable_quotes', [])
            if notable_quotes:
                doc.add_heading('Key Excerpts', 2)
                for nq in notable_quotes[:10]:
                    if isinstance(nq, dict):
                        quote = nq.get('quote', '')
                        significance = nq.get('significance', '')
                        if quote:
                            doc.add_paragraph(f'"{quote}"')
                        if significance:
                            doc.add_paragraph(f'    Significance: {significance}')
                    else:
                        doc.add_paragraph(f'"{nq}"')

            # Recommendations
            recommendations = ai_analysis.get('recommendations', [])
            if recommendations:
                doc.add_heading('Recommendations', 2)
                for rec in recommendations:
                    doc.add_paragraph(f'  {rec}')

            doc.add_page_break()

        # === Legal Team Summary ===
        if legal_summary:
            doc.add_heading('Legal Team Summary', 1)
            doc.add_paragraph(
                'This section provides a comprehensive narrative summary of the analysis '
                'results, written for the legal team. It explains the key findings and '
                'how to use the accompanying output files.'
            )
            markdown_to_docx(doc, legal_summary)
            doc.add_page_break()

        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(self._generate_executive_summary(
            extracted_data, analysis_results, review_decisions
        ))
        
        # Data Overview
        doc.add_heading('Data Overview', 1)

        # Calculate metadata from extracted_data structure
        messages = extracted_data.get('messages', extracted_data.get('combined', []))
        total_messages = len(messages) if isinstance(messages, list) else 0

        date_range = self._compute_date_range(messages)
        sources = set()
        if messages and total_messages > 0:
            for msg in messages:
                if msg.get('source'):
                    sources.add(msg['source'])

        screenshots = extracted_data.get('screenshots', [])
        threats = analysis_results.get('threats', {})
        threat_summary = threats.get('summary', {})
        threat_details = threats.get('details', [])
        messages_with_threats = threat_summary.get('messages_with_threats', 0)

        overview_rows = [
            ('Metric', 'Value'),
            ('Total Messages', str(total_messages)),
            ('Date Range', date_range),
            ('Sources', ', '.join(sources) if sources else 'N/A'),
            ('Threats Detected', str(messages_with_threats)),
            ('Items Reviewed', str(review_decisions.get('total_reviewed', 0))),
        ]
        if screenshots:
            overview_rows.append(('Screenshots Cataloged', str(len(screenshots))))
        overview_table = doc.add_table(rows=len(overview_rows), cols=2)
        for row_idx, (metric, value) in enumerate(overview_rows):
            overview_table.rows[row_idx].cells[0].text = metric
            overview_table.rows[row_idx].cells[1].text = value
        overview_table.columns[0].width = Inches(3.0)
        overview_table.columns[1].width = Inches(3.5)
        self._style_docx_table(overview_table)
        doc.add_paragraph('')  # spacer

        # Threat Analysis
        doc.add_heading('Threat Analysis', 1)
        doc.add_paragraph(f"Threats detected: {messages_with_threats}")
        
        # Show high priority threats if available
        if threat_details and isinstance(threat_details, list):
            high_priority = [t for t in threat_details if t.get('threat_detected')][:5]
            if high_priority:
                doc.add_heading('High Priority Threats', 2)
                for threat in high_priority:
                    content = threat.get('content', '')[:200]
                    ts = self.compliance.convert_to_local(threat.get('timestamp'))
                    sender = threat.get('sender', '')
                    ts_display = f" [{ts}]" if ts else ''
                    sender_display = f" — {sender}" if sender else ''
                    doc.add_paragraph(f"• {content}{sender_display}{ts_display}")
        
        # Sentiment Analysis
        doc.add_heading('Sentiment Analysis', 1)
        sentiment = analysis_results.get('sentiment', [])
        
        # Calculate sentiment distribution if we have data
        if sentiment and isinstance(sentiment, list):
            positive = sum(1 for s in sentiment if s.get('sentiment_polarity') == 'positive')
            negative = sum(1 for s in sentiment if s.get('sentiment_polarity') == 'negative')
            neutral = sum(1 for s in sentiment if s.get('sentiment_polarity') == 'neutral')

            doc.add_paragraph(f"Sentiment distribution:")
            doc.add_paragraph(f"  • Positive: {positive}")
            doc.add_paragraph(f"  • Neutral: {neutral}")
            doc.add_paragraph(f"  • Negative: {negative}")
        else:
            doc.add_paragraph("Sentiment analysis data not available")

        # Emotional Escalation Patterns from AI
        if ai_analysis:
            sentiment_ai = ai_analysis.get('sentiment_analysis', {})
            shifts = sentiment_ai.get('shifts', [])
            if shifts:
                doc.add_heading('Emotional Escalation Patterns', 2)
                doc.add_paragraph(
                    'The following emotional shifts were detected during pre-review screening, '
                    'indicating potential escalation patterns:'
                )
                for shift in shifts:
                    if isinstance(shift, dict):
                        from_state = shift.get('from', 'unknown')
                        to_state = shift.get('to', 'unknown')
                        position = shift.get('approximate_position', '')
                        doc.add_paragraph(f'    {from_state} -> {to_state} ({position})')
                    else:
                        doc.add_paragraph(f'    {shift}')

        # Manual Review Summary
        doc.add_heading('Manual Review', 1)
        doc.add_paragraph(f"Items reviewed: {review_decisions.get('total_reviewed', 0)}")
        doc.add_paragraph(f"Relevant: {review_decisions.get('relevant', 0)}")
        doc.add_paragraph(f"Not relevant: {review_decisions.get('not_relevant', 0)}")
        doc.add_paragraph(f"Uncertain: {review_decisions.get('uncertain', 0)}")

        # Third-Party Contacts
        third_party = extracted_data.get('third_party_contacts', [])
        if third_party:
            doc.add_heading('Third-Party Contacts', 1)
            doc.add_paragraph(
                f'{len(third_party)} third-party contacts were discovered during analysis '
                'from emails and screenshots. These are contacts not included in the '
                'configured person mappings.'
            )
            tp_rows = [('Identifier', 'Display Name', 'Source')]
            for entry in third_party:
                tp_rows.append((
                    entry.get('identifier', ''),
                    entry.get('display_name', ''),
                    ', '.join(entry.get('sources', [])),
                ))
            tp_table = doc.add_table(rows=len(tp_rows), cols=3)
            for row_idx, (ident, name, src) in enumerate(tp_rows):
                tp_table.rows[row_idx].cells[0].text = ident
                tp_table.rows[row_idx].cells[1].text = name
                tp_table.rows[row_idx].cells[2].text = src
            tp_table.columns[0].width = Inches(2.5)
            tp_table.columns[1].width = Inches(2.0)
            tp_table.columns[2].width = Inches(2.0)
            self._style_docx_table(tp_table)

        # Chain of Custody
        doc.add_heading('Chain of Custody', 1)
        doc.add_paragraph(
            f"Total recorded actions: {len(self.forensic.actions)}"
        )
        doc.add_paragraph(
            f"Session ID: {self.forensic.session_id}"
        )
        doc.add_paragraph(
            f"Session start: {self.compliance.convert_to_local(self.forensic.start_time)}"
        )
        doc.add_paragraph('See accompanying chain_of_custody.json for the detailed forensic trail.')

        # Save document
        output_path = self.output_dir / f"forensic_report_{timestamp}.docx"
        doc.save(output_path)
        
        # Record hash
        file_hash = self.forensic.compute_hash(output_path)
        self.forensic.record_action(
            "word_report_generated",
            f"Generated Word report with hash {file_hash}",
            {"path": str(output_path), "hash": file_hash}
        )
        
        return output_path

    def _generate_json_report(self, extracted_data: Dict, analysis_results: Dict,
                            review_decisions: Dict, timestamp: str,
                            legal_summary: str = None) -> Path:
        """Generate JSON report."""
        report = {
            "metadata": {
                "type": "Forensic Message Analysis Report",
                "generated": datetime.now().isoformat(),
                "case_id": timestamp,
                "version": "1.0"
            },
            "extraction": extracted_data,
            "analysis": analysis_results,
            "review": review_decisions,
            "legal_team_summary": legal_summary,
            "summary": {
                "total_messages": len(extracted_data.get('messages', [])),
                "threats_detected": analysis_results.get('threats', {}).get('summary', {}).get('messages_with_threats', 0),
                "items_reviewed": review_decisions.get('total_reviewed', 0),
                "relevant_items": review_decisions.get('relevant', 0)
            }
        }
        
        output_path = self.output_dir / f"forensic_report_{timestamp}.json"
        
        with open(output_path, 'w') as f:
            json.dump(report, f, indent=2, default=str)
        
        # Record hash
        file_hash = self.forensic.compute_hash(output_path)
        self.forensic.record_action(
            "json_report_generated",
            f"Generated JSON report with hash {file_hash}",
            {"path": str(output_path), "hash": file_hash}
        )

        return output_path

    def _generate_legal_summary_docx(self, legal_summary: str, output_path: Path,
                                      reports: Dict[str, Any] = None):
        """Generate a formatted Word document from the legal team summary text.

        Parses the narrative and produces a professional document with case header, formatted paragraphs, an output file reference table, and a compliance footer.

        Args:
            legal_summary: Plain text narrative.
            output_path: Path for the output .docx file.
            reports: Dict mapping report type keys to file paths. Used to build the output file reference table with actual filenames.
        """
        doc = Document()

        # Default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Title
        title = doc.add_heading('Legal Team Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Case header
        header = self.compliance.generate_report_header()
        case_numbers = header.get('case_numbers') or [header['case_number']]
        if len(case_numbers) > 1:
            doc.add_paragraph('Case Numbers:')
            for cn in case_numbers:
                doc.add_paragraph(f'  • {cn}')
        else:
            doc.add_paragraph(f"Case Number: {case_numbers[0]}")
        doc.add_paragraph(f"Case Name: {header['case_name']}")
        doc.add_paragraph(f"Generated: {header['date_of_examination']}")
        if header['examiner_name'] != 'Not specified':
            doc.add_paragraph(f"Examiner: {header['examiner_name']}")
        doc.add_paragraph('')  # spacer

        # Body -- parse markdown formatting
        markdown_to_docx(doc, legal_summary)

        # Output file reference table
        if reports:
            doc.add_paragraph('')
            doc.add_heading('Output File Reference', level=1)
            doc.add_paragraph(
                'The following files were generated alongside this summary. '
                'All files are in the same output directory.'
            )

            # Define descriptions and guidance for each report type
            file_info = {
                'excel': (
                    'Excel Report',
                    'Start here. Contains per-person tabs with messages, integrated '
                    'threat/sentiment columns, a Findings Summary sheet, Timeline, '
                    'Conversation Threads, and Third Party Contacts. '
                    'Use filters and search to locate specific conversations.'
                ),
                'word': (
                    'Word Report',
                    'Comprehensive narrative report with case information, '
                    'findings summary, threat analysis, sentiment analysis, '
                    'and chain of custody reference. Suitable for court filing.'
                ),
                'methodology': (
                    'Methodology Statement',
                    'Standalone document explaining, in plain language, '
                    'every step of the analysis pipeline and how each '
                    'FRE / Daubert standard was satisfied. Read this '
                    'first if anyone questions the methodology; it is '
                    'written for judges and attorneys, not technicians.'
                ),
                'pdf': (
                    'PDF Report',
                    'Same content as the Word report, formatted for distribution '
                    'and printing. Use this for court submission or sharing with '
                    'opposing counsel.'
                ),
                'html': (
                    'HTML Report',
                    'Interactive report with per-person message tables, inline images, '
                    'conversation threads, risk indicators, and legal appendices '
                    '(Methodology, Completeness Validation, Limitations). Open in any '
                    'web browser.'
                ),
                'html_pdf': (
                    'HTML Report (PDF)',
                    'PDF conversion of the HTML report for printing or court submission.'
                ),
                'chat_html': (
                    'Chat Bubble Report',
                    'iMessage-style conversation view with left/right aligned message '
                    'bubbles, inline attachments, edit history, deleted message badges, '
                    'and URL previews. Best for reviewing conversations in context.'
                ),
                'json': (
                    'JSON Report',
                    'Machine-readable output with all analysis data. Intended for '
                    'technical review or import into other forensic tools.'
                ),
                'forensic_json': (
                    'Forensic JSON Report',
                    'Complete forensic data export with metadata and chain of custody.'
                ),
                'timeline': (
                    'Interactive Timeline',
                    'Chronological visualization of flagged events (threats, SOS, '
                    'patterns) alongside all email communications. Open in any web '
                    'browser. Use this to see the full chronology of events.'
                ),
                'chain_of_custody': (
                    'Chain of Custody',
                    'Complete forensic audit trail with SHA-256 hashes, timestamps, '
                    'and all operations performed. Required for FRE 901 authentication.'
                ),
                'manifest': (
                    'Run Manifest',
                    'Documentation of all inputs processed, outputs generated, and '
                    'processing steps taken. Supports FRE 803(6) business records.'
                ),
            }

            table = doc.add_table(rows=1, cols=3)

            # Header row
            hdr = table.rows[0].cells
            for i, text in enumerate(['File', 'Type', 'How to Use']):
                hdr[i].text = text

            for key, path_val in reports.items():
                path = Path(str(path_val))
                filename = path.name

                if key in file_info:
                    label, guidance = file_info[key]
                else:
                    label = key.replace('_', ' ').title()
                    guidance = ''

                row = table.add_row().cells
                # Filename in bold blue
                fn_para = row[0].paragraphs[0]
                fn_run = fn_para.add_run(filename)
                fn_run.bold = True
                fn_run.font.size = Pt(9)
                fn_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

                # Type label
                type_para = row[1].paragraphs[0]
                type_run = type_para.add_run(label)
                type_run.font.size = Pt(10)

                # Guidance
                guide_para = row[2].paragraphs[0]
                guide_run = guide_para.add_run(guidance)
                guide_run.font.size = Pt(10)

            # Set column widths and apply consistent blue theme
            for row in table.rows:
                row.cells[0].width = Inches(2.5)
                row.cells[1].width = Inches(1.3)
                row.cells[2].width = Inches(3.7)
            self._style_docx_table(table)

            # Note about files generated after this summary
            note_para = doc.add_paragraph()
            note_run = note_para.add_run(
                'Note: The interactive timeline, chain of custody, and run manifest '
                'are generated after this summary and will also be present in the '
                'output directory.'
            )
            note_run.font.size = Pt(9)
            note_run.font.italic = True

        # Compliance footer
        doc.add_paragraph('')
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(
            'This summary was generated by the Forensic Message Analyzer '
            f"v{header['tools_used'].split('v')[-1] if 'v' in header['tools_used'] else 'N/A'} "
            'using AI-assisted analysis. Findings are supplementary and should be '
            'validated against the underlying evidence and accompanying forensic reports.'
        )
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True

        doc.save(str(output_path))

    def _legal_summary_report_rows(self, reports: Dict[str, Any]) -> list:
        """Build [(filename, type_label, guidance), ...] rows for the legal-summary report table in either format."""
        file_info = {
            'excel': ('Excel Report', 'Start here. Per-person tabs, findings summary, timeline, conversation threads, manual-review decisions, and third-party contacts.'),
            'word': ('Word Report', 'Comprehensive narrative report with case information, findings summary, threat analysis, sentiment analysis, and chain-of-custody reference. Suitable for court filing.'),
            'methodology': ('Methodology Statement', 'Standalone document explaining every step of the pipeline and how each FRE / Daubert factor was satisfied. Read this first if anyone questions the methodology.'),
            'methodology_pdf': ('Methodology PDF', 'PDF version of the Methodology Statement for court exhibits.'),
            'pdf': ('PDF Report', 'Same content as the Word report, formatted for distribution and printing. Use this for court submission.'),
            'html': ('HTML Report', 'HTML with inline attachment images; legal appendices for Methodology, Completeness Validation, Limitations.'),
            'html_pdf': ('HTML → PDF', 'PDF rendering of the HTML report (via WeasyPrint).'),
            'chat': ('Chat Report', 'iMessage-style chat-bubble transcript; open in a browser.'),
            'chat_html': ('Chat HTML', 'iMessage-style chat-bubble transcript; open in a browser.'),
            'events_timeline': ('Events Timeline', 'Sparse court-facing chronology of the moments the case turns on.'),
            'timeline': ('Detailed Timeline', 'Minute-level chronological view for analyst drill-down.'),
            'json': ('JSON Report', 'Machine-readable raw analysis output.'),
        }
        rows = []
        for key, path in reports.items():
            if key == 'legal_summary':
                continue
            filename = Path(str(path)).name
            label, guidance = file_info.get(key, (key.replace('_', ' ').title(), ''))
            rows.append([filename, label, guidance])
        return rows

    def _generate_legal_team_summary(self, extracted_data: Dict,
                                     analysis_results: Dict,
                                     review_decisions: Dict) -> str:
        """
        Generate a comprehensive narrative summary for the legal team using Claude.

        Explains all findings and how to interpret the output files.
        Returns None if AI is not available.
        """
        try:
            from anthropic import Anthropic
        except ImportError:
            logger.info("Anthropic not available, skipping legal team summary")
            return None

        if not self.config.ai_api_key:
            logger.info("AI API key not configured, skipping legal team summary")
            return None

        # Skip API call if there's no actual data to summarize
        messages = extracted_data.get('messages', extracted_data.get('combined', []))
        total_messages = len(messages) if isinstance(messages, list) else 0
        if total_messages == 0:
            logger.info("No messages to summarize, skipping legal team summary")
            return None

        # Source breakdown
        source_counts = {}
        if messages and isinstance(messages, list):
            for msg in messages:
                src = msg.get('source', 'unknown')
                source_counts[src] = source_counts.get(src, 0) + 1

        # Date range
        date_range = self._compute_date_range(messages)

        # Threat stats
        threats = analysis_results.get('threats', {})
        threat_count = threats.get('summary', {}).get('messages_with_threats', 0)
        threat_categories = threats.get('summary', {}).get('category_counts', {})

        # Sentiment stats
        sentiment = analysis_results.get('sentiment', [])
        sentiment_dist = {'positive': 0, 'neutral': 0, 'negative': 0}
        if sentiment and isinstance(sentiment, list):
            for s in sentiment:
                pol = s.get('sentiment_polarity', 'neutral')
                if pol in sentiment_dist:
                    sentiment_dist[pol] += 1

        # Pre-review screening stats
        ai_analysis = analysis_results.get('ai_analysis', {})
        risk_indicators = ai_analysis.get('risk_indicators', [])
        ai_threat_severity = ai_analysis.get('threat_assessment', {}).get('severity', 'none')
        recommendations = ai_analysis.get('recommendations', [])

        # Review stats
        total_reviewed = review_decisions.get('total_reviewed', 0)
        relevant = review_decisions.get('relevant', 0)

        # Third-party contacts
        third_party = extracted_data.get('third_party_contacts', [])

        # Build the prompt
        context = (
            f"DATASET OVERVIEW:\n"
            f"- Total messages analyzed: {total_messages}\n"
            f"- Date range: {date_range}\n"
            f"- Sources: {', '.join(f'{src}: {count}' for src, count in source_counts.items())}\n"
            f"- Screenshots cataloged: {len(extracted_data.get('screenshots', []))}\n\n"
            f"THREAT ANALYSIS:\n"
            f"- Messages with threats detected: {threat_count}\n"
            f"- Threat categories: {json.dumps(threat_categories) if threat_categories else 'None'}\n\n"
            f"AI RISK ASSESSMENT:\n"
            f"- Overall threat severity: {ai_threat_severity}\n"
            f"- Risk indicators found: {len(risk_indicators)}\n"
        )
        for ri in risk_indicators[:10]:
            if isinstance(ri, dict):
                sev = ri.get('severity', 'unknown')
                desc = ri.get('indicator', ri.get('description', ri.get('detail', '')))
                context += f"  [{sev}] {desc}\n"
            else:
                context += f"  {ri}\n"

        context += (
            f"\nSENTIMENT ANALYSIS:\n"
            f"- Positive messages: {sentiment_dist['positive']}\n"
            f"- Neutral messages: {sentiment_dist['neutral']}\n"
            f"- Negative messages: {sentiment_dist['negative']}\n\n"
            f"MANUAL REVIEW:\n"
            f"- Items reviewed: {total_reviewed}\n"
            f"- Confirmed relevant: {relevant}\n\n"
            f"THIRD-PARTY CONTACTS:\n"
            f"- Unmapped contacts discovered: {len(third_party)}\n\n"
            f"AI RECOMMENDATIONS:\n"
        )
        for rec in recommendations[:10]:
            context += f"- {rec}\n"

        context += (
            f"\nOUTPUT FILES GENERATED:\n"
            f"- Excel report (.xlsx): Contains per-person tabs with filtered messages, "
            f"threat indicators, and sentiment data for each configured party. "
            f"Start here for exploring specific conversations.\n"
            f"- Word report (.docx): Narrative analysis with findings summary, "
            f"risk indicators, threat details, and chain of custody\n"
            f"- PDF report (.pdf): Same content as Word, formatted for court submission\n"
            f"- HTML report (.html): Interactive report with per-person message tables, "
            f"inline images, conversation threads, and legal appendices\n"
            f"- Chat-bubble report (_chat.html): iMessage-style conversation view with "
            f"message bubbles, edit history, deleted message badges, and URL previews. "
            f"Best for reading conversations in context.\n"
            f"- Timeline (.html): Interactive chronological visualization of all messages "
            f"and flagged events\n"
            f"- Chain of custody (.json): Complete forensic audit trail\n"
            f"- Run manifest (.json): Documentation of all inputs, outputs, and processing steps\n"
            f"\nA detailed file reference table with exact filenames and usage guidance "
            f"will be appended to the legal team summary document after this narrative.\n"
        )

        system_prompt = (
            "You are a forensic analyst writing a comprehensive summary for "
            "attorneys and paralegals working on a family law matter. Your summary "
            "should be written in plain, professional language suitable for legal "
            "professionals who are NOT technicians.\n\n"
            "Write a narrative summary that:\n"
            "1. Opens with a brief overview of what was analyzed and the scope of the data\n"
            "2. Summarizes the most important findings — threats, risks, and concerning patterns\n"
            "3. Explains the sentiment analysis results and what they mean for the case\n"
            "4. Notes any third-party contacts discovered that may be relevant\n"
            "5. Describes what each output file contains and how the legal team should use it:\n"
            "   - Which file to open first\n"
            "   - How to find specific conversations in the Excel report\n"
            "   - How to interpret threat and sentiment columns\n"
            "   - When to reference the timeline vs. the Excel report\n"
            "6. Closes with recommended next steps for the legal team\n\n"
            "Keep it to 4-6 paragraphs. Use factual language. Do not speculate beyond "
            "what the data shows. Reference specific numbers from the analysis."
        )

        try:
            client = Anthropic(
                api_key=self.config.ai_api_key,
                base_url="https://api.anthropic.com",
            )
            model = (
                self.config.ai_summary_model
                or self.config.ai_batch_model
                or 'claude-sonnet-4-6'
            )
            response = client.messages.create(
                model=model,
                system=[{
                    "type": "text",
                    "text": system_prompt,
                    "cache_control": {"type": "ephemeral"},
                }],
                messages=[{"role": "user", "content": context}],
                max_tokens=2048,
            )
            result = response.content[0].text

            # Track tokens from this sync API call (standard rates, not batch)
            legal_input = response.usage.input_tokens
            legal_output = response.usage.output_tokens

            rp = get_pricing(model)
            sync_cost = (legal_input / 1_000_000) * rp['input'] + (legal_output / 1_000_000) * rp['output']

            # Update AI processing stats if available
            ai_stats = analysis_results.get('ai_analysis', {}).get('processing_stats')
            if ai_stats is not None:
                ai_stats["input_tokens"] = ai_stats.get("input_tokens", 0) + legal_input
                ai_stats["output_tokens"] = ai_stats.get("output_tokens", 0) + legal_output
                ai_stats["tokens_used"] = ai_stats.get("tokens_used", 0) + legal_input + legal_output
                ai_stats["api_calls"] = ai_stats.get("api_calls", 0) + 1
                ai_stats["legal_summary_sync_cost_usd"] = round(sync_cost, 4)
                ai_stats["estimated_cost_usd"] = round(
                    ai_stats.get("estimated_cost_usd", 0) + sync_cost, 4
                )

            print(f"    Legal team summary: {legal_input:,} input + {legal_output:,} output tokens (~${sync_cost:.4f})")

            self.forensic.record_action(
                "legal_team_summary_generated",
                "Generated AI-powered legal team summary",
                {"model": model, "length": len(result),
                 "input_tokens": legal_input, "output_tokens": legal_output,
                 "estimated_cost_usd": round(sync_cost, 4)}
            )
            return result

        except Exception as e:
            logger.error(f"Failed to generate legal team summary: {e}")
            self.forensic.record_action(
                "legal_team_summary_error",
                f"Error generating legal team summary: {str(e)}",
                {"error": str(e)}
            )
            return None

    def _generate_executive_summary(self, extracted_data: Dict,
                                   analysis_results: Dict,
                                   review_decisions: Dict) -> str:
        """Generate executive summary for legal team review."""
        messages = extracted_data.get('messages', extracted_data.get('combined', []))
        total_messages = len(messages) if isinstance(messages, list) else 0
        threats = analysis_results.get('threats', {}).get('summary', {}).get('messages_with_threats', 0)
        reviewed = review_decisions.get('total_reviewed', 0)
        relevant = review_decisions.get('relevant', 0)

        ai_analysis = analysis_results.get('ai_analysis', {})
        ai_summary = ai_analysis.get('conversation_summary', '')
        risk_count = len(ai_analysis.get('risk_indicators', []))
        ai_threats_found = ai_analysis.get('threat_assessment', {}).get('found', False)
        ai_threat_severity = ai_analysis.get('threat_assessment', {}).get('severity', 'none')

        summary = (
            f"This forensic analysis examined {total_messages} digital communications "
            f"extracted from multiple sources. Automated screening flagged "
            f"{threats} messages containing potentially threatening or concerning content "
            f"for manual review. "
        )

        if reviewed > 0:
            summary += (
                f"Of the items flagged for manual review, {reviewed} were examined by "
                f"a qualified analyst, with {relevant} confirmed as relevant to the proceedings. "
            )

        if ai_summary and 'not available' not in ai_summary.lower() and 'not configured' not in ai_summary.lower():
            summary += (
                f"\n\nAdditional automated screening identified {risk_count} distinct risk "
                f"indicators warranting attention. "
            )
            if ai_threats_found:
                summary += (
                    f"Overall assessed severity: {ai_threat_severity}. "
                    f"All such items were submitted to the same manual review process. "
                )
            summary += f"\n\n{ai_summary}"

        summary += (
            "\n\nAll data handling maintained forensic integrity through SHA-256 "
            "cryptographic hashing and comprehensive chain of custody documentation. "
            "See the Findings Summary section for detailed risk indicators, key excerpts, "
            "and recommended actions."
        )

        return summary.strip()

    def _generate_limitations(self, analysis_results: Dict) -> list:
        """Generate limitation statements based on available data and features."""
        return generate_limitations(self.config, analysis_results)
